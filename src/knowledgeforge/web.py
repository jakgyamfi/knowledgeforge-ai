"""FastAPI application: capture, transcription, organization, projects, and chat."""

from __future__ import annotations
import logging
import mimetypes
import re
import threading
from contextlib import asynccontextmanager
from datetime import datetime, timezone
from pathlib import Path
from typing import Annotated, Literal
from uuid import uuid4
from fastapi import BackgroundTasks, FastAPI, File, HTTPException, Query, UploadFile
from fastapi.responses import FileResponse, Response
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel, Field, SecretStr
from .ai import AIService
from .config import Settings
from .library import Library
from .organizer import organize_transcript
from .pipeline import AUDIO_EXTENSIONS, TranscriptionPipeline


class NoteUpdate(BaseModel):
    title: str = Field(min_length=1, max_length=200)
    category: str = Field(default="author", max_length=50)
    tags: list[str] = Field(default_factory=list, max_length=30)
    summary: str = Field(default="", max_length=100_000)
    project_id: int | None = None


class ProjectCreate(BaseModel):
    name: str = Field(min_length=1, max_length=120)
    description: str = Field(default="", max_length=2000)
    workspace_type: str = Field(default="book", pattern="^(book|business|impact|project|general)$")


class ProjectUpdate(BaseModel):
    name: str = Field(min_length=1, max_length=120)
    description: str = Field(default="", max_length=10_000)


class ProfileUpdate(BaseModel):
    display_name: str = Field(default="", max_length=200)
    location: str = Field(default="", max_length=200)
    headline: str = Field(default="", max_length=500)
    summary: str = Field(default="", max_length=20_000)
    skills: list[str] = Field(default_factory=list, max_length=300)
    certifications: list[str] = Field(default_factory=list, max_length=100)
    interests: list[str] = Field(default_factory=list, max_length=200)
    goals: list[str] = Field(default_factory=list, max_length=100)
    industries: list[str] = Field(default_factory=list, max_length=100)
    avoid: list[str] = Field(default_factory=list, max_length=100)
    preferences: dict = Field(default_factory=dict)


class OpportunityAction(BaseModel):
    status: str = Field(pattern="^(new|saved|dismissed)$")


class WorkspaceCardCreate(BaseModel):
    collection: str = Field(min_length=1, max_length=80, pattern="^[a-z0-9_-]+$")
    title: str = Field(min_length=1, max_length=300)
    content: str = Field(default="", max_length=50_000)
    metadata: dict = Field(default_factory=dict)


class WorkspaceCardUpdate(BaseModel):
    collection: str = Field(min_length=1, max_length=80, pattern="^[a-z0-9_-]+$")
    title: str = Field(min_length=1, max_length=300)
    content: str = Field(default="", max_length=50_000)


class CompletionConfirm(BaseModel):
    confirmed: bool


class ChatRequest(BaseModel):
    question: str = Field(min_length=2, max_length=5000)
    project_id: int | None = None
    selected_note_ids: list[int] = Field(default_factory=list, max_length=100)


class AISelection(BaseModel):
    provider: str = Field(min_length=1, max_length=80, pattern="^[a-z0-9_-]+$")
    model: str = Field(min_length=1, max_length=120)


class ProviderSecret(BaseModel):
    api_key: SecretStr


class BookInstructions(BaseModel):
    instructions: str = Field(default="", max_length=50_000)


class ReorganizeRequest(BaseModel):
    feedback: str = Field(default="", max_length=50_000)


class WorkspaceResearchRequest(BaseModel):
    query: str = Field(min_length=3, max_length=2000)


class WorkspaceImproveRequest(BaseModel):
    selection: str = Field(min_length=3, max_length=50_000)
    instruction: str = Field(default="Make this clearer, stronger, and more actionable.", max_length=5000)


class WorkspaceAccelerateRequest(BaseModel):
    objective: str = Field(default="", max_length=10_000)


class TaskCreate(BaseModel):
    title: str = Field(min_length=1, max_length=300)
    details: str = Field(default="", max_length=5000)
    priority: str = Field(default="medium", pattern="^(critical|high|medium|low)$")
    estimate_minutes: int = Field(default=30, ge=5, le=10080)
    target_date: str = Field(default="", max_length=10)


class TaskUpdate(BaseModel):
    title: str | None = Field(default=None, min_length=1, max_length=300)
    details: str | None = Field(default=None, max_length=5000)
    status: str | None = Field(default=None, pattern="^(todo|doing|done|dismissed)$")
    priority: str | None = Field(default=None, pattern="^(critical|high|medium|low)$")
    estimate_minutes: int | None = Field(default=None, ge=5, le=10080)
    target_date: str | None = Field(default=None, max_length=10)
    position: int | None = Field(default=None, ge=0)


class GrowthItemUpdate(BaseModel):
    status: str | None = Field(default=None, pattern="^(planned|active|in_progress|paused|completed)$")
    progress: int | None = Field(default=None, ge=0, le=100)
    target_date: str | None = Field(default=None, max_length=10)
    notes: str | None = Field(default=None, max_length=5000)


class GrowthActionUpdate(BaseModel):
    status: str | None = Field(default=None, pattern="^(todo|doing|done|dismissed)$")
    title: str | None = Field(default=None, min_length=1, max_length=300)
    details: str | None = Field(default=None, max_length=5000)
    priority: str | None = Field(default=None, pattern="^(critical|high|medium|low)$")
    estimate_minutes: int | None = Field(default=None, ge=5, le=10080)
    target_date: str | None = Field(default=None, max_length=10)


def _safe_filename(filename: str) -> str:
    name = Path(filename or "recording.webm").name
    return re.sub(r"[^A-Za-z0-9._ -]+", "-", name).strip(" .-") or "recording.webm"


def create_app() -> FastAPI:
    settings = Settings.load()
    library = Library(settings.database)
    library.ensure_setting("ai_provider", settings.ai_provider)
    library.ensure_setting("ai_model", settings.ai_model)
    library.ensure_setting("auto_integrate", "true")
    library.ensure_project(settings.default_project, "Default workspace for captured book ideas.", "book")
    library.ensure_project("Business Ideas", "Captured business concepts and working plans.", "business")
    library.ensure_project("Project Ideas", "Captured project concepts and implementation plans.", "project")
    library.sync_growth_from_profile()
    pipeline = TranscriptionPipeline(settings, library)
    ai = AIService(settings, library)
    stop_event = threading.Event()
    worker = threading.Thread(target=pipeline.watch, args=(stop_event,), name="knowledgeforge-watcher", daemon=True)
    analysis_lock = threading.Lock()
    analysis_queue = {"running": False, "last": None}

    def profile_view(profile: dict) -> dict:
        """Return profile metadata to the local UI without echoing extracted CV text."""
        result = dict(profile)
        result["cv_configured"] = bool(result.pop("cv_text", ""))
        result.pop("cv_path", None)
        return result

    def process_pending_analysis() -> dict[str, int | bool | str | None]:
        """Drain unfinished analysis once without overlapping another drain.

        This function is intentionally event-driven rather than a tight retry
        loop. That prevents a bad key or provider outage from repeatedly
        consuming tokens. Startup, key save, model selection, and the recovery
        button are the supported retry events.
        """
        if not analysis_lock.acquire(blocking=False):
            return {"accepted": False, "processed": 0, "failed": 0, "message": "Analysis queue is already running."}
        analysis_queue["running"] = True
        processed = failed = 0
        try:
            if not ai.enabled:
                return {
                    "accepted": False,
                    "processed": 0,
                    "failed": 0,
                    "message": "The selected AI provider is not configured.",
                }
            for note_id in library.pending_analysis_ids():
                try:
                    organized = ai.analyze(note_id)
                    project_id = organized.get("project_id")
                    if project_id and library.get_setting("auto_integrate", "true") == "true":
                        ai.reorganize_book(project_id, trigger_note_id=note_id)
                    processed += 1
                except Exception as exc:
                    failed += 1
                    library.mark_analysis_failed(note_id, str(exc))
                    logging.exception("Queued AI analysis failed for note %s", note_id)
            return {"accepted": True, "processed": processed, "failed": failed, "message": "Queue pass complete."}
        finally:
            analysis_queue["running"] = False
            analysis_queue["last"] = datetime.now(timezone.utc).isoformat()
            analysis_lock.release()

    def start_pending_analysis() -> None:
        """Run a queue pass in a daemon so API requests return immediately."""
        if analysis_queue["running"] or not ai.enabled or not library.pending_analysis_ids(limit=1):
            return
        threading.Thread(
            target=process_pending_analysis,
            name="knowledgeforge-analysis-queue",
            daemon=True,
        ).start()

    def process_import(path: Path) -> None:
        """Extract imported documents/images, then run the same organization workflow as audio."""
        suffix = path.suffix.lower()
        if suffix in {".txt", ".md"}:
            content = path.read_text(encoding="utf-8", errors="replace")
        elif suffix == ".pdf":
            from pypdf import PdfReader

            content = "\n\n".join(page.extract_text() or "" for page in PdfReader(path).pages)
        elif suffix == ".docx":
            from docx import Document

            content = "\n".join(paragraph.text for paragraph in Document(path).paragraphs)
        elif suffix in {".png", ".jpg", ".jpeg", ".webp", ".gif"}:
            content = ai.describe_image(path)
        else:
            raise ValueError(f"Unsupported import type: {suffix}")
        transcript_path = settings.transcripts / "imports" / f"{path.stem}.txt"
        transcript_path.parent.mkdir(parents=True, exist_ok=True)
        transcript_path.write_text(content.strip() + "\n", encoding="utf-8")
        note_id = library.upsert_transcript(
            source_name=f"imports/{path.name}",
            audio_path=path,
            transcript_path=transcript_path,
            transcript=content.strip(),
        )
        if ai.enabled:
            organized = ai.analyze(note_id)
            if organized.get("project_id") and library.get_setting("auto_integrate", "true") == "true":
                ai.reorganize_book(organized["project_id"], trigger_note_id=note_id)

    def extract_profile_document(path: Path) -> str:
        suffix = path.suffix.lower()
        if suffix in {".txt", ".md"}:
            return path.read_text(encoding="utf-8", errors="replace")
        if suffix == ".pdf":
            from pypdf import PdfReader

            return "\n\n".join(page.extract_text() or "" for page in PdfReader(path).pages)
        if suffix == ".docx":
            from docx import Document

            return "\n".join(paragraph.text for paragraph in Document(path).paragraphs)
        raise ValueError("Profile documents must be TXT, Markdown, PDF, or DOCX.")

    @asynccontextmanager
    async def lifespan(_: FastAPI):
        worker.start()
        start_pending_analysis()
        yield
        stop_event.set()
        worker.join(timeout=max(settings.poll_seconds + 2, 10))

    app = FastAPI(
        title="KnowledgeForge AI",
        version="0.11.0",
        description=(
            "Security-first AI idea orchestration for developing private, unstructured source material "
            "into books, business opportunities, and projects."
        ),
        lifespan=lifespan,
    )
    static_dir = Path(__file__).parent / "static"
    app.mount("/assets", StaticFiles(directory=static_dir), name="assets")

    @app.get("/", include_in_schema=False)
    def home():
        return FileResponse(static_dir / "index.html")

    @app.get("/health")
    def health():
        return {
            "status": "ok",
            "library": library.stats(),
            "watcher_alive": worker.is_alive(),
            "ai_enabled": ai.enabled,
            "ai_provider": ai.provider,
            "ai_model": ai.model,
            "analysis_queue": {
                "running": analysis_queue["running"],
                "last": analysis_queue["last"],
            },
        }

    @app.get("/api/ai/options")
    def ai_options():
        return ai.options()

    @app.put("/api/ai/selection")
    def select_ai(payload: AISelection, background: BackgroundTasks):
        options = ai.options()
        provider = next((item for item in options["providers"] if item["id"] == payload.provider), None)
        if provider is None:
            raise HTTPException(422, "Unknown provider; add it to config/ai-providers.json")
        # OpenAI choices are curated. Ollama names come from the local service;
        # a custom Ollama model is accepted when the service is temporarily off.
        if provider["id"] != "ollama" and payload.model not in provider["models"]:
            raise HTTPException(422, "Unsupported catalog model")
        library.set_settings({"ai_provider": payload.provider, "ai_model": payload.model.strip()})
        logging.info("AI selection changed to %s / %s", payload.provider, payload.model)
        background.add_task(start_pending_analysis)
        return ai.options()

    @app.get("/api/secrets/status")
    def secret_status():
        return ai.secret_status()

    @app.put("/api/secrets/{provider_id}")
    def save_provider_secret(provider_id: str, payload: ProviderSecret, background: BackgroundTasks):
        if settings.web_host not in {"127.0.0.1", "localhost", "::1"}:
            raise HTTPException(403, "In-app secret entry is allowed only on a loopback-bound desktop instance.")
        try:
            status = ai.set_secret(provider_id, payload.api_key.get_secret_value())
            # If this key belongs to the active provider, old notes created
            # before configuration are now eligible for automatic analysis.
            background.add_task(start_pending_analysis)
            return status
        except KeyError as exc:
            raise HTTPException(404, "Provider does not accept an API key") from exc
        except RuntimeError as exc:
            raise HTTPException(409, str(exc)) from exc

    @app.delete("/api/secrets/{provider_id}")
    def delete_provider_secret(provider_id: str):
        if settings.web_host not in {"127.0.0.1", "localhost", "::1"}:
            raise HTTPException(403, "In-app secret management is allowed only on a loopback-bound desktop instance.")
        try:
            return ai.delete_secret(provider_id)
        except KeyError as exc:
            raise HTTPException(404, "Provider does not accept an API key") from exc
        except RuntimeError as exc:
            raise HTTPException(409, str(exc)) from exc

    @app.get("/api/books/{project_id}")
    def get_book(project_id: int):
        try:
            book = library.get_book(project_id)
        except KeyError as exc:
            raise HTTPException(404, "Book project not found") from exc
        book["revisions"] = library.list_revisions(project_id)
        return book

    @app.put("/api/books/{project_id}/instructions")
    def save_instructions(project_id: int, payload: BookInstructions):
        try:
            library.get_book(project_id)
        except KeyError as exc:
            raise HTTPException(404, "Book project not found") from exc
        library.save_book_instructions(project_id, payload.instructions)
        return library.get_book(project_id)

    @app.post("/api/books/{project_id}/reorganize")
    def reorganize_book(project_id: int, payload: ReorganizeRequest):
        try:
            return ai.reorganize_book(project_id, feedback=payload.feedback)
        except RuntimeError as exc:
            raise HTTPException(503, str(exc)) from exc
        except Exception as exc:
            logging.exception("Book reorganization failed")
            raise HTTPException(502, "Book reorganization failed; see private log.") from exc

    @app.get("/api/projects/{project_id}/tasks")
    def project_tasks(project_id: int):
        try:
            library.get_book(project_id)
        except KeyError as exc:
            raise HTTPException(404, "Workspace not found") from exc
        return library.list_tasks(project_id)

    @app.post("/api/projects/{project_id}/tasks", status_code=201)
    def create_task(project_id: int, payload: TaskCreate):
        try:
            library.get_book(project_id)
        except KeyError as exc:
            raise HTTPException(404, "Workspace not found") from exc
        return library.add_task(project_id, **payload.model_dump())

    @app.patch("/api/projects/{project_id}/tasks/{task_id}")
    def update_task(project_id: int, task_id: int, payload: TaskUpdate):
        task = library.update_task(project_id, task_id, payload.model_dump(exclude_unset=True))
        if task is None:
            raise HTTPException(404, "Task not found")
        return task

    @app.post("/api/projects/{project_id}/tasks/replan")
    def replan_tasks(project_id: int):
        try:
            return ai.refresh_plan(project_id, reason="Owner requested a fresh execution plan")
        except KeyError as exc:
            raise HTTPException(404, "Workspace not found") from exc
        except RuntimeError as exc:
            raise HTTPException(503, str(exc)) from exc
        except Exception as exc:
            logging.exception("Execution-plan refresh failed")
            raise HTTPException(502, "Execution-plan refresh failed; see private log.") from exc

    @app.get("/api/books/{project_id}/export")
    def export_book(project_id: int):
        try:
            book = library.get_book(project_id)
        except KeyError as exc:
            raise HTTPException(404, "Book project not found") from exc
        content = f"# {book['project']['name']}\n\n" + "\n\n".join(
            f"## {section['title']}\n\n{section['content']}" for section in book["sections"]
        )
        filename = re.sub(r"[^A-Za-z0-9._-]+", "-", book["project"]["name"]).strip("-") or "book"
        return Response(
            content,
            media_type="text/markdown",
            headers={"Content-Disposition": f'attachment; filename="{filename}.md"'},
        )

    @app.get("/api/projects")
    def projects():
        return library.list_projects()

    @app.post("/api/projects", status_code=201)
    def create_project(payload: ProjectCreate):
        project_id = library.ensure_project(payload.name, payload.description, payload.workspace_type)
        return next(p for p in library.list_projects() if p["id"] == project_id)

    @app.get("/api/opportunities")
    def opportunities(status: str = Query(default="new", pattern="^(new|saved|validated|exploring|completed|dismissed)$")):
        return library.list_opportunities(status)

    @app.post("/api/opportunities/refresh")
    def refresh_opportunities():
        try:
            return ai.refresh_opportunities()
        except RuntimeError as exc:
            raise HTTPException(503, str(exc)) from exc
        except Exception as exc:
            logging.exception("Opportunity refresh failed")
            raise HTTPException(502, "Opportunity refresh failed; see private log.") from exc

    @app.patch("/api/opportunities/{opportunity_id}")
    def update_opportunity(opportunity_id: int, payload: OpportunityAction):
        item = library.update_opportunity(opportunity_id, status=payload.status)
        if not item:
            raise HTTPException(404, "Opportunity not found")
        return item

    @app.post("/api/opportunities/{opportunity_id}/validate")
    def validate_opportunity(opportunity_id: int):
        try:
            return ai.validate_opportunity(opportunity_id)
        except KeyError as exc:
            raise HTTPException(404, "Opportunity not found") from exc
        except Exception as exc:
            logging.exception("Opportunity validation failed")
            raise HTTPException(502, "Internet validation failed; see private log.") from exc

    @app.post("/api/opportunities/{opportunity_id}/explore")
    def explore_opportunity(opportunity_id: int):
        opportunity = library.get_opportunity(opportunity_id)
        result = library.explore_opportunity(opportunity_id)
        if result is None:
            raise HTTPException(404, "Opportunity not found")
        try:
            return ai.initialize_workspace(result["project_id"], opportunity)
        except Exception as exc:
            logging.exception("Workspace initialization failed")
            raise HTTPException(502, "Workspace creation failed; see private log.") from exc

    @app.get("/api/profile")
    def profile():
        return profile_view(library.get_profile())

    @app.get("/manual")
    def user_manual_markdown():
        manual = Path(__file__).parent / "manuals" / "USER-MANUAL.md"
        if not manual.is_file():
            raise HTTPException(404, "User manual is not installed")
        return FileResponse(manual, filename="KnowledgeForge-User-Manual.md", media_type="text/markdown")

    @app.get("/manual.pdf")
    def user_manual_pdf():
        manual = Path(__file__).parent / "manuals" / "KnowledgeForge-User-Manual.pdf"
        if not manual.is_file():
            raise HTTPException(404, "PDF user manual is not installed")
        return FileResponse(manual, filename=manual.name, media_type="application/pdf")

    @app.put("/api/profile")
    def update_profile(payload: ProfileUpdate):
        result = profile_view(library.update_profile(payload.model_dump()))
        library.sync_growth_from_profile()
        return result

    @app.delete("/api/profile/suggestions")
    def clear_profile_suggestions():
        return profile_view(library.clear_profile_suggestions())

    @app.post("/api/profile/suggest")
    def suggest_profile():
        try:
            return ai.suggest_profile_updates()
        except RuntimeError as exc:
            raise HTTPException(503, str(exc)) from exc
        except Exception as exc:
            logging.exception("Profile suggestion failed")
            raise HTTPException(502, "Profile suggestion failed; see private log.") from exc

    @app.get("/api/growth")
    def growth_overview():
        return library.growth_overview()

    @app.post("/api/growth/plan")
    def refresh_growth_plan(weeks: Literal[2, 4, 12] = 4):
        try:
            return ai.refresh_growth_plan(weeks)
        except RuntimeError as exc:
            raise HTTPException(503, str(exc)) from exc
        except Exception as exc:
            logging.exception("Growth planning failed")
            raise HTTPException(502, "Growth planning failed; see private log.") from exc

    @app.patch("/api/growth/items/{item_id}")
    def update_growth_item(item_id: int, payload: GrowthItemUpdate):
        item = library.update_growth_item(item_id, payload.model_dump(exclude_unset=True))
        if item is None:
            raise HTTPException(404, "Growth item not found")
        return item

    @app.patch("/api/growth/actions/{action_id}")
    def update_growth_action(action_id: int, payload: GrowthActionUpdate):
        action = library.update_growth_action(action_id, payload.model_dump(exclude_unset=True))
        if action is None:
            raise HTTPException(404, "Growth action not found")
        return action

    @app.post("/api/profile/cv")
    async def upload_profile_cv(file: Annotated[UploadFile, File()]):
        destination = settings.imports / "profile" / f"{uuid4().hex}-{_safe_filename(file.filename or 'cv')}"
        destination.parent.mkdir(parents=True, exist_ok=True)
        payload = await file.read(settings.max_upload_mb * 1024 * 1024 + 1)
        if len(payload) > settings.max_upload_mb * 1024 * 1024:
            raise HTTPException(413, "File exceeds configured upload limit")
        destination.write_bytes(payload)
        try:
            text = extract_profile_document(destination)
        except ValueError as exc:
            destination.unlink(missing_ok=True)
            raise HTTPException(415, str(exc)) from exc
        return profile_view(
            library.update_profile(
                {
                    "cv_text": text[:200_000],
                    "cv_filename": file.filename or destination.name,
                    "cv_path": str(destination),
                }
            )
        )

    def resolve_profile_cv() -> Path:
        """Resolve only the CV path recorded in the private profile."""
        profile = library.get_profile()
        stored = Path(profile.get("cv_path", ""))
        profile_root = (settings.imports / "profile").resolve()
        try:
            candidate = stored.resolve()
            candidate.relative_to(profile_root)
        except (OSError, ValueError):
            raise HTTPException(404, "No stored CV is available") from None
        if not candidate.is_file():
            raise HTTPException(404, "The stored CV file is missing")
        return candidate

    @app.get("/api/profile/cv")
    def download_profile_cv():
        destination = resolve_profile_cv()
        profile = library.get_profile()
        return FileResponse(
            destination,
            filename=profile.get("cv_filename") or destination.name,
            media_type=mimetypes.guess_type(destination.name)[0] or "application/octet-stream",
        )

    @app.delete("/api/profile/cv")
    def remove_profile_cv():
        destination = resolve_profile_cv()
        destination.unlink(missing_ok=True)
        return profile_view(library.update_profile({"cv_text": "", "cv_filename": "", "cv_path": ""}))

    @app.get("/api/workspaces/{project_id}")
    def workspace(project_id: int):
        try:
            return library.workspace_snapshot(project_id)
        except KeyError as exc:
            raise HTTPException(404, "Workspace not found") from exc

    @app.patch("/api/projects/{project_id}")
    def update_project(project_id: int, payload: ProjectUpdate):
        try:
            project = library.update_project(project_id, **payload.model_dump())
        except ValueError as exc:
            raise HTTPException(409, str(exc)) from exc
        if project is None:
            raise HTTPException(404, "Workspace not found")
        return project

    @app.delete("/api/projects/{project_id}", status_code=204)
    def delete_project(project_id: int):
        if not library.delete_project(project_id):
            raise HTTPException(404, "Workspace not found")
        return Response(status_code=204)

    @app.post("/api/workspaces/{project_id}/initialize")
    def initialize_workspace(project_id: int):
        try:
            return ai.initialize_workspace(project_id)
        except KeyError as exc:
            raise HTTPException(404, "Workspace not found") from exc
        except Exception as exc:
            logging.exception("Workspace initialization failed")
            raise HTTPException(502, "Workspace initialization failed; see private log.") from exc

    @app.post("/api/workspaces/{project_id}/cards", status_code=201)
    def add_workspace_card(project_id: int, payload: WorkspaceCardCreate):
        try:
            library.get_book(project_id)
        except KeyError as exc:
            raise HTTPException(404, "Workspace not found") from exc
        return library.add_workspace_card(project_id, **payload.model_dump())

    @app.patch("/api/workspaces/{project_id}/cards/{card_id}")
    def update_workspace_card(project_id: int, card_id: int, payload: WorkspaceCardUpdate):
        card = library.update_workspace_card(project_id, card_id, **payload.model_dump())
        if card is None:
            raise HTTPException(404, "Workspace card not found")
        return card

    @app.post("/api/workspaces/{project_id}/research")
    def research_workspace(project_id: int, payload: WorkspaceResearchRequest):
        try:
            return ai.research_workspace(project_id, payload.query)
        except KeyError as exc:
            raise HTTPException(404, "Workspace not found") from exc
        except RuntimeError as exc:
            raise HTTPException(503, str(exc)) from exc
        except Exception as exc:
            logging.exception("Workspace research failed")
            raise HTTPException(502, "Workspace research failed; see private log.") from exc

    @app.post("/api/workspaces/{project_id}/improve")
    def improve_workspace_selection(project_id: int, payload: WorkspaceImproveRequest):
        try:
            return ai.improve_workspace_selection(project_id, payload.selection, payload.instruction)
        except KeyError as exc:
            raise HTTPException(404, "Workspace not found") from exc
        except RuntimeError as exc:
            raise HTTPException(503, str(exc)) from exc
        except Exception as exc:
            logging.exception("Workspace improvement failed")
            raise HTTPException(502, "Workspace improvement failed; see private log.") from exc

    @app.post("/api/workspaces/{project_id}/accelerate")
    def accelerate_workspace(project_id: int, payload: WorkspaceAccelerateRequest):
        try:
            return ai.accelerate_workspace(project_id, payload.objective)
        except KeyError as exc:
            raise HTTPException(404, "Workspace not found") from exc
        except RuntimeError as exc:
            raise HTTPException(503, str(exc)) from exc
        except Exception as exc:
            logging.exception("Workspace acceleration failed")
            raise HTTPException(502, "Workspace acceleration failed; see private log.") from exc

    @app.post("/api/workspaces/{project_id}/complete")
    def complete_workspace(project_id: int, payload: CompletionConfirm):
        if not payload.confirmed:
            raise HTTPException(422, "Completion must be explicitly confirmed.")
        try:
            return ai.complete_workspace(project_id)
        except KeyError as exc:
            raise HTTPException(404, "Workspace not found") from exc
        except RuntimeError as exc:
            raise HTTPException(409, str(exc)) from exc
        except Exception as exc:
            logging.exception("Workspace completion review failed")
            raise HTTPException(502, "Workspace completion failed; see private log.") from exc

    @app.post("/api/workspaces/{project_id}/reopen")
    def reopen_workspace(project_id: int):
        try:
            return library.reopen_workspace(project_id)
        except KeyError as exc:
            raise HTTPException(404, "Workspace not found") from exc

    @app.get("/api/notes")
    def notes(
        q: str = Query(default="", max_length=200),
        category: str = Query(default="", max_length=50),
        project_id: int | None = None,
    ):
        return library.list_notes(q.strip(), category.strip(), project_id)

    @app.get("/api/notes/{note_id}")
    def note(note_id: int):
        item = library.get_note(note_id)
        if not item:
            raise HTTPException(404, "Note not found")
        return item

    @app.patch("/api/notes/{note_id}")
    def update_note(note_id: int, payload: NoteUpdate):
        tags = sorted({tag.strip().lower() for tag in payload.tags if tag.strip()})
        if not library.update_note(
            note_id,
            title=payload.title,
            category=payload.category,
            tags=tags,
            summary=payload.summary,
            project_id=payload.project_id,
        ):
            raise HTTPException(404, "Note not found")
        return library.get_note(note_id)

    @app.post("/api/notes/{note_id}/analyze")
    def analyze(note_id: int):
        if not library.get_note(note_id):
            raise HTTPException(404, "Note not found")
        try:
            organized = ai.analyze(note_id)
            if organized.get("project_id") and library.get_setting("auto_integrate", "true") == "true":
                ai.reorganize_book(organized["project_id"], trigger_note_id=note_id)
            return library.get_note(note_id)
        except RuntimeError as exc:
            raise HTTPException(503, str(exc)) from exc
        except Exception as exc:
            library.mark_analysis_failed(note_id, str(exc))
            logging.exception("Analysis failed")
            raise HTTPException(502, "AI analysis failed; details are in the private log.") from exc

    @app.post("/api/analysis/process-pending", status_code=202)
    def process_pending(background: BackgroundTasks):
        if not ai.enabled:
            raise HTTPException(409, "Configure and select an AI provider first.")
        background.add_task(start_pending_analysis)
        return {
            "accepted": True,
            "pending": library.stats()["pending_analysis"],
            "running": analysis_queue["running"],
        }

    @app.post("/api/notes/{note_id}/organize")
    def local_organize(note_id: int):
        item = library.get_note(note_id)
        if not item:
            raise HTTPException(404, "Note not found")
        draft = organize_transcript(item["transcript"])
        library.update_note(
            note_id,
            title=item["title"],
            category=item["category"],
            tags=list(draft["tags"]),
            summary=str(draft["summary"]),
            project_id=item["project_id"],
        )
        return library.get_note(note_id)

    @app.post("/api/chat")
    def chat(payload: ChatRequest):
        try:
            return ai.answer(payload.question, payload.project_id, payload.selected_note_ids)
        except RuntimeError as exc:
            raise HTTPException(503, str(exc)) from exc
        except Exception as exc:
            logging.exception("Library chat failed")
            raise HTTPException(502, "AI chat failed; see private log.") from exc

    @app.post("/api/scan", status_code=202)
    def scan(background: BackgroundTasks):
        background.add_task(pipeline.run_once)
        return {"accepted": True}

    @app.post("/api/upload", status_code=202)
    async def upload(file: Annotated[UploadFile, File()]):
        original = _safe_filename(file.filename or "recording.webm")
        extension = Path(original).suffix.lower()
        if extension not in AUDIO_EXTENSIONS:
            raise HTTPException(415, f"Unsupported audio type: {extension or 'none'}")
        destination = settings.inbox / f"{uuid4().hex[:8]}-{original}"
        maximum = settings.max_upload_mb * 1024 * 1024
        written = 0
        try:
            with destination.open("wb") as output:
                while chunk := await file.read(1024 * 1024):
                    written += len(chunk)
                    if written > maximum:
                        raise HTTPException(413, f"Upload exceeds {settings.max_upload_mb} MB")
                    output.write(chunk)
        except Exception:
            destination.unlink(missing_ok=True)
            raise
        finally:
            await file.close()
        logging.info("Accepted upload: %s", destination.name)
        return {"accepted": True, "filename": destination.name}

    @app.post("/api/import", status_code=202)
    async def import_content(background: BackgroundTasks, file: Annotated[UploadFile, File()]):
        original = _safe_filename(file.filename or "content.txt")
        allowed = {".txt", ".md", ".pdf", ".docx", ".png", ".jpg", ".jpeg", ".webp", ".gif"}
        if Path(original).suffix.lower() not in allowed:
            raise HTTPException(415, "Supported imports: TXT, Markdown, PDF, DOCX, PNG, JPG, WEBP, GIF")
        destination = settings.imports / f"{uuid4().hex[:8]}-{original}"
        maximum = settings.max_upload_mb * 1024 * 1024
        written = 0
        try:
            with destination.open("wb") as output:
                while chunk := await file.read(1024 * 1024):
                    written += len(chunk)
                    if written > maximum:
                        raise HTTPException(413, f"Upload exceeds {settings.max_upload_mb} MB")
                    output.write(chunk)
        except Exception:
            destination.unlink(missing_ok=True)
            raise
        finally:
            await file.close()
        background.add_task(process_import, destination)
        return {"accepted": True, "filename": destination.name}

    @app.get("/api/notes/{note_id}/audio")
    def audio(note_id: int):
        item = library.get_note(note_id)
        if not item:
            raise HTTPException(404, "Note not found")
        path = Path(item["audio_path"])
        if not path.exists():
            raise HTTPException(404, "Archived audio is unavailable")
        return FileResponse(path, media_type=mimetypes.guess_type(path.name)[0] or "application/octet-stream")

    @app.delete("/api/notes/{note_id}/audio")
    def delete_audio(note_id: int):
        """Delete only KnowledgeForge-managed source media, never arbitrary files."""
        item = library.get_note(note_id)
        if not item:
            raise HTTPException(404, "Note not found")
        if not item["audio_path"]:
            return {"deleted": False, "message": "Source audio was already removed."}
        path = Path(item["audio_path"]).resolve()
        managed_roots = (settings.recordings.resolve(), settings.imports.resolve())
        if not any(path.is_relative_to(root) for root in managed_roots):
            raise HTTPException(409, "This source is outside KnowledgeForge-managed storage and was not deleted.")
        if path.exists() and path.is_file():
            path.unlink()
        library.clear_note_audio(note_id)
        return {"deleted": True, "transcript_preserved": True}

    return app
